"""DOCX converter via python-docx (optional dep). Walks paragraphs +
tables, prefixes Word heading styles with markdown # so the semantic
chunker can later split on them."""

from __future__ import annotations

from pathlib import Path

from . import ArtifactConversionError


def read_docx(p: Path) -> str:
    try:
        import docx  # type: ignore[import-not-found]
    except ImportError as exc:
        raise ArtifactConversionError(
            "DOCX support requires python-docx. Install with one of:\n"
            "    pip install proofagent-harness[artifact]\n"
            "    pip install python-docx"
        ) from exc

    try:
        d = docx.Document(str(p))
    except Exception as exc:
        raise ArtifactConversionError(
            f"Failed to open DOCX {p}: {type(exc).__name__}: {exc}"
        ) from exc

    parts: list[str] = []
    for para in d.paragraphs:
        text = (para.text or "").strip()
        if not text:
            continue
        style = (getattr(para.style, "name", "") or "").lower()
        if style.startswith("heading 1"):
            parts.append(f"# {text}")
        elif style.startswith("heading 2"):
            parts.append(f"## {text}")
        elif style.startswith("heading 3"):
            parts.append(f"### {text}")
        else:
            parts.append(text)

    for tbl in d.tables:
        rows: list[str] = []
        for row in tbl.rows:
            cells = [(c.text or "").strip().replace("\n", " ") for c in row.cells]
            rows.append("| " + " | ".join(cells) + " |")
        if rows:
            if len(rows) > 1:
                sep = "| " + " | ".join(["---"] * len(tbl.rows[0].cells)) + " |"
                rows.insert(1, sep)
            parts.append("\n".join(rows))

    return "\n\n".join(parts) if parts else f"[DOCX appeared empty: {p.name}]"
